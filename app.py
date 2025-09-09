import streamlit as st
import pandas as pd
import plotly.express as px
import json, re, base64, hashlib, secrets, logging, zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict
from io import BytesIO
from mimetypes import guess_type
from decimal import Decimal, InvalidOperation

# ================================
# TCHAI â€” Easy LCA Indicator (v4+)
# -------------------------------
# âœ“ Sign-in (3 pre-created users)
# âœ“ Settings â†’ Upload & Activate a PERMANENT database (persists until changed)
# âœ“ Inputs: tolerant parsing, NO Excel previews, clear process dropdowns
# âœ“ Workspace: Results & Comparison â†’ Final Summary â†’ Report (PDF) â†’ Versions
# âœ“ User Guide: first page after sign-in; inline if possible, always downloadable
# âœ“ PDF Report: smart-filled from live inputs; DOCX fallback if PDF backend missing
# âœ“ Safe folder creation (avoids FileExistsError)
# âœ“ Caching for Excel + parsing (step 5)
# âœ“ Better number parsing (step 6)
# âœ“ Session schema w/ Pydantic (step 7)
# âœ“ Smarter column detection (step 8)
# âœ“ Better status/tooltips (step 11)
# âœ“ Simple i18n hooks (step 13)
# âœ“ Robust DOCX replacements (step 16)
# âœ“ Sanitize version names (step 18)
# âœ“ Logging (step 21)
# âœ“ Pinned requirements (step 22)
# âœ“ Font fallback embedding (step 23)
# ================================

st.set_page_config(
    page_title="TCHAI â€” Easy LCA Indicator",
    page_icon="ðŸŒ¿",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -----------------------------
# Safe dirs (resolve relative to CWD)
# -----------------------------

def ensure_dir(p: Path):
    if p.exists() and not p.is_dir():
        backup = p.with_name(f"{p.name}.conflict.{datetime.now().strftime('%Y%m%d%H%M%S')}")
        p.rename(backup)
    p.mkdir(parents=True, exist_ok=True)

APP_DIR = Path.cwd()
ASSETS  = APP_DIR / "assets";      ensure_dir(ASSETS)
DB_ROOT = ASSETS / "databases";    ensure_dir(DB_ROOT)
GUIDES  = ASSETS / "guides";       ensure_dir(GUIDES)
FONTS   = ASSETS / "fonts";        ensure_dir(FONTS)
USERS_FILE     = ASSETS / "users.json"
ACTIVE_DB_FILE = DB_ROOT / "active.json"  # stores {"path": "...xlsx"}

# -----------------------------
# Logging (step 21)
# -----------------------------
LOGS_DIR = ASSETS / "logs"; ensure_dir(LOGS_DIR)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    handlers=[
        logging.FileHandler(LOGS_DIR / "app.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger("tchai")

# -----------------------------
# Optional PDF/DOCX backends
# -----------------------------
REPORTLAB_OK = False
DOCX_OK = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    REPORTLAB_OK = True
except Exception:
    REPORTLAB_OK = False

try:
    from docx import Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# -----------------------------
# Branding (logo)
# -----------------------------
LOGO_CANDIDATES = [ASSETS / "tchai_logo.png", Path("tchai_logo.png"), Path("/mnt/data/tchai_logo.png")]
_logo_bytes = None
for p in LOGO_CANDIDATES:
    if p.exists():
        try:
            _logo_bytes = p.read_bytes()
            break
        except Exception:
            pass

def logo_tag(height=86):
    if not _logo_bytes:
        return "<span style='font-weight:900;font-size:28px'>TCHAI</span>"
    b64 = base64.b64encode(_logo_bytes).decode()
    return f"<img src='data:image/png;base64,{b64}' alt='TCHAI' style='height:{height}px'/>"

# -----------------------------
# i18n (step 13)
# -----------------------------
LANG_FILE_DIR = ASSETS / "i18n"; ensure_dir(LANG_FILE_DIR)
if "lang" not in st.session_state: st.session_state.lang = "en"

def t(key: str, default: Optional[str]=None) -> str:
    path = LANG_FILE_DIR / f"{st.session_state.lang}.json"
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data.get(key, default or key)
    except Exception:
        return default or key

# -----------------------------
# Visual Theme (Light Oat + PP Neue Montreal + Pop color) + font fallback (step 23)
# -----------------------------
BG = "#E4E5DA"       # Light Oat
POP = "#B485FF"      # Pop color for comparison charts

def _embed_font_css() -> str:
    font_regular = FONTS / "PPNeueMontreal-Regular.woff2"
    font_medium  = FONTS / "PPNeueMontreal-Medium.woff2"
    def face(path: Path, weight: int) -> str:
        try:
            if path.exists():
                b64 = base64.b64encode(path.read_bytes()).decode()
                return f"""
                @font-face {{
                  font-family: 'PP Neue Montreal';
                  src: url(data:font/woff2;base64,{b64}) format('woff2');
                  font-weight: {weight};
                  font-style: normal;
                  font-display: swap;
                }}
                """
        except Exception:
            logger.exception("Font embed failed")
        return ""
    return face(font_regular, 400) + face(font_medium, 500)

font_css = _embed_font_css()
st.markdown(
    f"""
    <style>
      {font_css}
      /* Fallback to file URLs if data-URI missing */
      @font-face {{
        font-family: 'PP Neue Montreal';
        src: url('assets/fonts/PPNeueMontreal-Regular.woff2') format('woff2'),
             url('assets/fonts/PPNeueMontreal-Regular.ttf') format('truetype');
        font-weight: 400;
        font-style: normal;
        font-display: swap;
      }}
      @font-face {{
        font-family: 'PP Neue Montreal';
        src: url('assets/fonts/PPNeueMontreal-Medium.woff2') format('woff2'),
             url('assets/fonts/PPNeueMontreal-Medium.ttf') format('truetype');
        font-weight: 500;
        font-style: normal;
        font-display: swap;
      }}

      .stApp {{
        background: {BG};
        color: #000;
        font-family: 'PP Neue Montreal', ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, "Apple Color Emoji","Segoe UI Emoji";
      }}

      html, body, [class*="css"] {{ font-weight: 400; }}

      h1, h2, h3, h4, .brand-title, .stTabs [data-baseweb="tab"] p {{
        font-weight: 500 !important;
        text-transform: capitalize;
        letter-spacing: 0.2px;
      }}

      .stRadio > label, .stRadio div[role="radiogroup"] label p {{
        font-weight: 500;
        text-transform: capitalize;
      }}

      .metric {{
        border: 1px solid #111;
        border-radius: 12px;
        padding: 14px;
        text-align: center;
        background: rgba(255,255,255,0.6);
        backdrop-filter: blur(2px);
      }}

      .brand-title {{ font-size: 26px; text-align: center; }}

      .stSelectbox div[data-baseweb="select"],
      .stNumberInput input,
      .stTextInput input,
      .stTextArea textarea {{
        border: 1px solid #111;
        background: #fff;
      }}

      .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {{
        box-shadow: inset 0 -2px 0 0 {POP};
      }}
    </style>
    """,
    unsafe_allow_html=True,
)

# -----------------------------
# Utils
# -----------------------------

def _rerun():
    if hasattr(st, "rerun"): st.rerun()
    else:
        try: st.experimental_rerun()
        except Exception: pass

# -----------------------------
# Auth (simple salted SHA256 kept; Argon2 recommended later)
# -----------------------------

def _load_users() -> dict:
    try:
        return json.loads(USERS_FILE.read_text())
    except Exception:
        return {}

def _save_users(users: dict):
    USERS_FILE.write_text(json.dumps(users, indent=2))

def _hash(pw: str, salt: str) -> str:
    return hashlib.sha256((salt + pw).encode()).hexdigest()

def _initials(name: str) -> str:
    parts = [p for p in re.split(r"\s+|_+|\.+|@", name) if p]
    return ((parts[0][0] if parts else "U") + (parts[1][0] if len(parts) > 1 else "")).upper()

def bootstrap_users_if_needed():
    users = _load_users()
    if users: return
    default_pw = "ChangeMe123!"
    emails = ["sustainability@tchai.nl","jillderegt@tchai.nl","veravanbeaumont@tchai.nl"]
    out = {}
    for email in emails:
        salt = secrets.token_hex(8)
        out[email] = {"salt": salt, "hash": _hash(default_pw, salt), "created_at": datetime.now().isoformat()}
    _save_users(out)

bootstrap_users_if_needed()
if "auth_user" not in st.session_state:
    st.session_state.auth_user = None

# -----------------------------
# DB helpers
# -----------------------------

def list_databases() -> List[Path]:
    return sorted(DB_ROOT.glob("*.xlsx"), key=lambda p: p.stat().st_mtime, reverse=True)

def set_active_database(path: Path):
    ACTIVE_DB_FILE.write_text(json.dumps({"path": str(path)}))
    st.success(f"Activated database: {path.name}")
    _rerun()

def get_active_database_path() -> Optional[Path]:
    if ACTIVE_DB_FILE.exists():
        try:
            data = json.loads(ACTIVE_DB_FILE.read_text())
            p = Path(data.get("path", ""))
            if p.exists():
                return p
        except Exception:
            pass
    dbs = list_databases()
    if dbs:
        return dbs[0]
    for candidate in [ASSETS / "Refined database.xlsx", Path("Refined database.xlsx"), Path("database.xlsx")]:
        if candidate.exists():
            return candidate
    return None

# -----------------------------
# Caching helpers (step 5)
# -----------------------------
@st.cache_data(show_spinner=False)
def _open_xls_cached(path_str: str, mtime: float) -> pd.ExcelFile:
    return pd.ExcelFile(path_str)

def load_active_excel() -> Optional[pd.ExcelFile]:
    p = get_active_database_path()
    if p and p.exists():
        try:
            return _open_xls_cached(str(p), p.stat().st_mtime)
        except Exception as e:
            logger.exception("Active Excel open failed")
            st.error(f"Failed to open Excel: {p.name} â€” {e}")
            return None
    return None

@st.cache_data(show_spinner=False)
def _df_sig(df: pd.DataFrame) -> str:
    try:
        return hashlib.sha256(pd.util.hash_pandas_object(df.fillna(''), index=True).values).hexdigest()
    except Exception:
        # fallback to rows dump
        return hashlib.sha256(df.to_csv(index=False).encode()).hexdigest()

@st.cache_data(show_spinner=False)
def _parse_materials_cached(df: pd.DataFrame, sig: str) -> dict:
    return parse_materials(df)

@st.cache_data(show_spinner=False)
def _parse_processes_cached(df: pd.DataFrame, sig: str) -> dict:
    return parse_processes(df)

# -----------------------------
# Parsing helpers (steps 6 & 8)
# -----------------------------

def extract_number(v):
    try:
        if isinstance(v, (int, float)):
            return float(v)
        s = str(v).strip()
        s = s.replace('\u2212','-')   # minus sign â†’ hyphen
        s = s.replace(',', '.')       # European decimals
        m = re.search(r"[-+]?\d*\.?\d+(e[-+]?\d+)?", s, flags=re.I)
        if not m:
            return 0.0
        return float(Decimal(m.group()))
    except (InvalidOperation, ValueError, TypeError):
        return 0.0

def _normalize_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    def _canon(col: str) -> str:
        c = str(col)
        c = c.replace("coâ‚‚","co2").replace("â‚‚","2")
        c = re.sub(r"\s+", " ", c, flags=re.I).strip().lower()
        c = re.sub(r"\(.*?\)", "", c)  # remove unit hints
        return c
    df.columns = [_canon(c) for c in df.columns]
    return df

def _find_sheet(xls: pd.ExcelFile, target: str) -> Optional[str]:
    names = xls.sheet_names
    for n in names:
        if n == target: return n
    t = re.sub(r"\s+", "", target.lower())
    for n in names:
        if re.sub(r"\s+", "", n.lower()) == t: return n
    for n in names:
        if target.lower() in n.lower(): return n
    return None

def parse_materials(df_raw: pd.DataFrame) -> dict:
    if df_raw is None or df_raw.empty: return {}
    df = _normalize_cols(df_raw)
    def pick(aliases):
        for a in aliases:
            if a in df.columns: return a
        return None
    col_name = pick(["material name","material","name","material_name"])
    col_co2  = pick(["co2e per kg","co2e/kg","co2e","co2 per kg","co2","emission factor","co2e factor","co2 factor"])
    col_rc   = pick(["recycled content","recycled content %","recycled","recycled %","recycle %","recycled_pct"])
    col_eol  = pick(["eol","end of life","end-of-life"])
    col_life = pick(["lifetime","life","lifespan","lifetime years"])
    col_circ = pick(["circularity","circ","circularity level"])
    if not col_name or not col_co2: return {}
    out = {}
    for _, r in df.iterrows():
        name = (str(r[col_name]).strip() if pd.notna(r[col_name]) else "")
        if not name: continue
        out[name] = {
            "COâ‚‚e (kg)": extract_number(r[col_co2]) if pd.notna(r[col_co2]) else 0.0,
            "Recycled Content": extract_number(r[col_rc]) if col_rc and pd.notna(r.get(col_rc, None)) else 0.0,
            "EoL": str(r[col_eol]).strip() if col_eol and pd.notna(r.get(col_eol, None)) else "Unknown",
            "Lifetime": str(r[col_life]).strip() if col_life and pd.notna(r.get(col_life, None)) else "Unknown",
            "Circularity": str(r[col_circ]).strip() if col_circ and pd.notna(r.get(col_circ, None)) else "Unknown",
        }
    return out

def parse_processes(df_raw: pd.DataFrame) -> dict:
    if df_raw is None or df_raw.empty: return {}
    df = _normalize_cols(df_raw)
    def pick(aliases):
        for a in aliases:
            if a in df.columns: return a
        return None
    col_proc = pick(["process type","process_type","process","step","operation","process name","name"])
    col_co2  = pick(["co2e","co2e (kg)","co2","emission","factor","co2e factor","emission factor","emission factor kg"])
    col_unit = pick(["unit","uom","units","measure","measurement"])
    if not col_proc or not col_co2: return {}
    out = {}
    for _, r in df.iterrows():
        name = (str(r[col_proc]).strip() if pd.notna(r[col_proc]) else "")
        if not name: continue
        out[name] = {
            "COâ‚‚e": extract_number(r[col_co2]) if pd.notna(r[col_co2]) else 0.0,
            "Unit": str(r[col_unit]).strip() if col_unit and pd.notna(r.get(col_unit, None)) else "",
        }
    return out

# -----------------------------
# Session storage + schema (step 7)
# -----------------------------
from pydantic import BaseModel, Field, ValidationError

class ProcStep(BaseModel):
    process: str = ""
    amount: float = 0.0
    co2e_per_unit: float = 0.0
    unit: str = ""

class Assessment(BaseModel):
    lifetime_weeks: int = 52
    selected_materials: List[str] = Field(default_factory=list)
    material_masses: Dict[str, float] = Field(default_factory=dict)
    processing_data: Dict[str, List[ProcStep]] = Field(default_factory=dict)

if "materials" not in st.session_state: 
    st.session_state.materials = {}
if "processes" not in st.session_state: 
    st.session_state.processes = {}
if "assessment" not in st.session_state:
    st.session_state.assessment = Assessment().model_dump()


def _ensure_assessment_model():
    try:
        model = Assessment(**st.session_state.get("assessment", {}))
        st.session_state.assessment = model.model_dump()
    except ValidationError:
        st.session_state.assessment = Assessment().model_dump()

# normalize once at startup
_ensure_assessment_model()

# -----------------------------
# Sidebar (logo + language + nav)
# -----------------------------
with st.sidebar:
    st.markdown(
        f"<div style='display:flex;justify-content:center;margin-bottom:10px'>{logo_tag(64)}</div>",
        unsafe_allow_html=True
    )
    st.selectbox("Language", options=["en","nl"], index=["en","nl"].index(st.session_state.lang) if st.session_state.lang in ["en","nl"] else 0, key="lang")

    if st.session_state.auth_user:
        nav_labels = [
            t("nav.user_guide","User Guide"),
            t("nav.tool","Actual Tool"),
            t("nav.results","Results"),
            t("nav.versions","Version"),
            t("nav.settings","Administrational Settings"),
        ]
        nav_map = {label: label for label in nav_labels}
        if "nav" in st.session_state and st.session_state.nav not in nav_labels:
            st.session_state.nav = nav_labels[0]
        sel = st.radio(
            "Navigate",
            nav_labels,
            index=nav_labels.index(st.session_state.get("nav", nav_labels[0])),
            key="nav",
        )
        page = nav_map[sel]
    else:
        page = "Sign in"

# -----------------------------
# Header (logo, title, avatar)
# -----------------------------
cl, cm, cr = st.columns([0.18, 0.64, 0.18])
with cl:
    st.markdown(f"{logo_tag(86)}", unsafe_allow_html=True)
with cm:
    st.markdown("<div class='brand-title'>Easy LCA Indicator</div>", unsafe_allow_html=True)
with cr:
    if st.session_state.auth_user:
        initials = _initials(st.session_state.auth_user)
        if hasattr(st, "popover"):
            with st.popover(f"ðŸ‘¤ {initials}"):
                st.write(f"Signed in as **{st.session_state.auth_user}**")
                st.markdown("---")
                st.subheader("Account Settings.")
                with st.form("change_pw_form", clear_on_submit=True):
                    cur = st.text_input("Current password", type="password")
                    new = st.text_input("New password", type="password")
                    conf = st.text_input("Confirm new password", type="password")
                    submitted = st.form_submit_button("Change password")
                    if submitted:
                        users = _load_users()
                        rec = users.get(st.session_state.auth_user)
                        if not rec or _hash(cur, rec["salt"]) != rec["hash"]:
                            st.error("Current password is incorrect.")
                        elif not new or new != conf:
                            st.error("New passwords don't match.")
                        else:
                            salt = secrets.token_hex(8)
                            rec["salt"] = salt
                            rec["hash"] = _hash(new, salt)
                            users[st.session_state.auth_user] = rec
                            _save_users(users)
                            st.success("Password changed.")
                st.markdown("---")
                if st.button("Sign out"):
                    st.session_state.auth_user = None
                    _rerun()
        else:
            if st.button("Sign out"):
                st.session_state.auth_user = None
                _rerun()

# -----------------------------
# Sign-in gate
# -----------------------------
if not st.session_state.auth_user:
    st.markdown("### Sign In To Continue.")
    t1, t2 = st.columns([0.55, 0.45])
    with t1:
        st.markdown("Use your TCHAI account email and password.")
        with st.form("login_form"):
            u = st.text_input("Email", key="login_u", placeholder="you@tchai.nl")
            p = st.text_input("Password", type="password", key="login_p")
            submitted = st.form_submit_button("Sign in")
        if submitted:
            users = _load_users()
            rec = users.get(u)
            if not rec:
                st.error("Unknown user.")
            elif _hash(p, rec["salt"]) != rec["hash"]:
                st.error("Wrong password.")
            else:
                st.session_state.auth_user = u
                st.session_state.nav = t("nav.user_guide","User Guide")
                st.success("Welcome!")
                _rerun()
    with t2:
        st.markdown("#### Need Changes?")
        st.caption("User creation is disabled. Ask an admin to add a new account.")
    st.stop()

# ---- User Guide content ----

def guidelines_content() -> dict:
    sections = {
        "12 Must Haves": """
### Materials
- FSC/PEFC wood; sustainable MDF where feasible  
- Water-based paints / low-VOC adhesives  
- Avoid PVC and problematic plastics where possible  
- High recycled content where quality allows  

### Design & Build
- Modular & repairable assemblies (swap parts, keep the core)  
- Avoid mixed-material laminates that block recycling  
- Label parts for sorting (material codes)  
- Standardize repeat parts across programs  

### Process & Logistics
- Run LCA-Light before locking specs  
- Transport-efficient (flat-pack, stackable, lighter)  
- Source locally when it truly reduces impact  
- Plan end-of-life (reuse/recycle routes documented)  
""",
        "5 Golden Rules": """
- Think circular â€“ plan reuse, modularity and end-of-life from day one  
- Start light â€“ run LCA-Light in briefing/concept to steer choices early  
- Choose smart materials â€“ prefer recycled/recyclable, FSC/PEFC  
- Cut waste â€“ simplify parts/finishes/packaging; avoid over-engineering  
- Design for disassembly â€“ standard fasteners, clear material separation  
""",
        "AI and Sustainability": """
# **Conscious Compute: The Tchai AI Playbook**

### Using AI, Sustainably 
Yes, we use AI, even in sustainability work. Elephant in the room: AI isnâ€™t â€œfree.â€  
It burns energy. Think of AI as a super-fast helper who lives in a big factory far away.  
Every time we ask it to do something, the factory spins up computers and uses energy.  

So we do use AI, but we try to use it like a scalpel, not a sledgehammer â€” smart, not more.  

### How to use it
- When weâ€™re starting from zero, AI is great for a first rough draft, a list of options, or a quick comparison.  
- When we already know roughly what we want, AI helps tighten text, check a list, or spot gaps.  
- For images, we keep it tight: few versions as possible, clear prompts, no endless rerolls.  

We donâ€™t use AI just because itâ€™s shiny. Every extra prompt costs energy and time.  

### We can try to keep the footprint low by
- Asking fewer, better questions  
- Batching tasks (ask AI to do similar things all at once, e.g. â€œGive me 5 headline optionsâ€ instead of 5 prompts)  
- Choosing lighter tools for simple jobs  
- Reusing what you already generated  

### Data care 
Treat AI like a postcard: assume others could read it.  
- Donâ€™t paste contracts, prices, personal data, or unreleased designs  
- Use approved tools and privacy settings  
- If in doubt, donâ€™t uploadâ€”ask IT  

### The law bit 
- Discoverable like email: chat logs can be requested in legal cases  
- Third-party issue: client NDAs may forbid sharing with AI vendors  
- GDPR applies: minimise/anonymise personal data, never upload special-category data  
- Training & retention: some AI tools use prompts/outputs unless you turn it off  
- Where data lives: cloud tools may store data outside the EU, only use approved ones  

### Quick rules to use AI safe and conscious, before you hit â€œGenerateâ€
1. Be clear about what you need (one sentence)  
2. Start small (one prompt, one image set)  
3. Cap yourself (max 3 iterations)  
4. Reuse outputs; donâ€™t restart from scratch  
5. Stop if the value isnâ€™t improving  
6. Summarise, donâ€™t copy; describe the problem instead of pasting the doc  
7. Redact when necessary: names, prices, IDs, client references  
""",
        "Easy LCA Indicator Tool": """
LCA-Light is our fast sustainability check. It helps you to get early-stage insights and make comparisons in a project with a focus on materials, layouts, processes, and End of Life.  
It can help you make informed, data-based decisions already from the design process without having to run a full Life Cycle Assessment.  
""",
        "How does it work": """
### 1) **Set the lifespan**  
Tell the tool how long our solution will be in use (in weeks). Thatâ€™s the baseline for all the math.  

### 2) **Add materials & processes**  
Add each individual material of the design by simply:  
- a) Pick the material from the database.  
- b) Enter the total mass used (kg).  
- c) Add the process steps (what happens to it: cutting, coating, etc.).  
- d) Quick check of auto-filled facts  

When you pick a material, the tool shows: COâ‚‚e factor (kg/kg), recycled %, density, and default end-of-life (recycle/incinerate/landfill).  
If itâ€™s not the right variant (e.g., powder-coated vs raw), choose a better dataset or flag it for update by emailing **sustainability@tchai.nl** (or Jill).  

### 3) **Let the tool crunch**  
Based on what you entered, LCA-Light calculates:  
- a) Embodied carbon (COâ‚‚e from materials + processes).  
- b) Weighted recycled content (how much of your mass is recycled).  
- c) End-of-life split (recycle / reuse / landfill assumptions).  

### 4) **Read the summary**  
Youâ€™ll get a clean snapshot with:  
- a) Total recycled content % (of overall mass).  
- b) Total COâ‚‚e for the concept.  
- c) Tree-equivalent signal (how much COâ‚‚e the design represents over its lifespan).  
- d) A material-by-material end-of-life view.  

### 5) **Compare versions**  
Use the dashboard to compare options (materials, finishes, layouts) on the same indicators (recycled content, carbon).  
Save different versions, revisit them, or remove them as the concept evolves.  

### 6) **How to request a change**  
Send an email to **sustainability@tchai.nl** with:  
- What you need changed (data or code)  
- Why (project/client need)  
- Deadline  
- Any source files or references  
""",
        "When do you use it": """
| User | Stage of Use | Purpose | Inputs Required |
|---|---|---|---|
| Sales | Briefing / Proposal | Compare design directions for the client pitch | Materials, finishes, size/dimensions |
| Designers | Concept phase | Evaluate multiple layout / material options | Drawings, materials, and dimensions |
| Engineers | Pre-technical / detailing | Assess impact of different alternatives | Material specs, processes, lifespan, quantity estimates |
| Project Leaders | Client discussions / final decision | Share directionally whatâ€™s better and why | Final inputs (materials, processes, lifespan) |

### When **NOT** to Use It
- You want to make hard claims (the LCA-Light is an indicator tool)  
- You donâ€™t have the basics; when the minimum required data is not available (accuracy would be too low)  
- You need an official LCA for external purposes and official documentation (complete LCA report required)  
""",
    }
    return sections

# -----------------------------
# PAGE: User Guide
# -----------------------------
if page == t("nav.user_guide","User Guide"):
    st.header("User Guide")
    content = guidelines_content()
    tabs = st.tabs([
        "12 Must Haves",
        "5 Golden Rules",
        "AI and Sustainability",
        "Easy LCA Indicator Tool",
        "How does it work",
        "When do you use it",
    ])
    tab_names = [
        "12 Must Haves",
        "5 Golden Rules",
        "AI and Sustainability",
        "Easy LCA Indicator Tool",
        "How does it work",
        "When do you use it",
    ]
    for i, tn in enumerate(tab_names):
        with tabs[i]:
            st.subheader(tn)
            st.write(content.get(tn, ""))

# -----------------------------
# Administrational Settings
# -----------------------------
if page in (t("nav.settings","Administrational Settings"), "Settings"):
    st.subheader("Database Manager")
    st.caption("Upload your Excel ONCE. It becomes the active database until you change it here.")

    active = get_active_database_path()
    if active:
        st.success(f"Active database: **{active.name}**")
    else:
        st.warning("No active database set.")

    up = st.file_uploader("Upload Excel (.xlsx) And Activate.", type=["xlsx"], key="db_upload")
    if up is not None:
        try:
            ts = datetime.now().strftime("%Y%m%d-%H%M%S")
            dest = DB_ROOT / f"database_{ts}.xlsx"
            data = up.read()
            # quick ZIP magic check for XLSX integrity
            try:
                import io
                with zipfile.ZipFile(io.BytesIO(data)) as z: z.testzip()
            except Exception as e:
                raise ValueError("Uploaded file is not a valid .xlsx") from e
            dest.write_bytes(data)
            set_active_database(dest)
        except Exception as e:
            logger.exception("Upload failed")
            st.error(f"Upload failed: {e}")

    st.markdown("### Available Databases.")
    dbs = list_databases()
    if not dbs:
        st.info("No databases found. Upload one above.")
    else:
        for pth in dbs:
            cols = st.columns([0.6,0.2,0.2])
            with cols[0]:
                st.write(f"**{pth.name}**  ")
                st.caption(f"{datetime.fromtimestamp(pth.stat().st_mtime).strftime('%Y-%m-%d %H:%M')}")
            with cols[1]:
                if active and pth.exists() and active and pth.resolve() == active.resolve():
                    st.success("Active")
                else:
                    if st.button("Activate", key=f"act_{pth.name}"):
                        set_active_database(pth)
            with cols[2]:
                if active and pth.exists() and active and pth.resolve() == active.resolve():
                    st.caption("(can't delete active)")
                else:
                    if st.button("ðŸ—‘ï¸ Delete", key=f"rm_{pth.name}"):
                        try:
                            pth.unlink(missing_ok=True)
                            st.success("Deleted.")
                            _rerun()
                        except Exception as e:
                            st.error(f"Delete failed: {e}")

# -----------------------------
# Actual Tool (Inputs)
# -----------------------------
if page in (t("nav.tool","Actual Tool"), "Inputs"):
    active_path = get_active_database_path()
    st.subheader("Database Status")
    if active_path:
        st.success(f"Active database: **{active_path.name}**")
    else:
        st.error("No active database found. Go to Administrational Settings â†’ Database Manager.")

    st.caption("Optional: override for THIS session only")
    override = st.file_uploader("Session Override (.xlsx).", type=["xlsx"], key="override_db")

    if override is not None:
        try:
            xls = pd.ExcelFile(override)
            st.info("Using the uploaded session override.")
        except Exception as e:
            logger.exception("Override Excel open failed")
            st.error(f"Could not open the uploaded Excel: {e}")
            st.stop()
    else:
        xls = load_active_excel()

    if not xls:
        st.error("No Excel could be opened. Go to Administrational Settings or use the override above.")
        st.stop()

    auto_mat = _find_sheet(xls, "Materials") or xls.sheet_names[0]
    auto_proc = _find_sheet(xls, "Processes") or (xls.sheet_names[1] if len(xls.sheet_names)>1 else xls.sheet_names[0])

    c2, c3 = st.columns(2)
    with c2:
        mat_choice = st.selectbox("Materials Sheet.", options=xls.sheet_names,
                                  index=xls.sheet_names.index(auto_mat) if auto_mat in xls.sheet_names else 0)
    with c3:
        proc_choice = st.selectbox("Processes Sheet.", options=xls.sheet_names,
                                   index=xls.sheet_names.index(auto_proc) if auto_proc in xls.sheet_names else 0)

    with st.status("Parsing selected sheetsâ€¦", expanded=False) as status:
        try:
            mats_df = pd.read_excel(xls, sheet_name=mat_choice)
            procs_df = pd.read_excel(xls, sheet_name=proc_choice)
            st.session_state.materials = _parse_materials_cached(mats_df, _df_sig(mats_df))
            st.session_state.processes = _parse_processes_cached(procs_df, _df_sig(procs_df))
            status.update(label="Parsing complete âœ…", state="complete")
        except Exception as e:
            status.update(label=f"Parsing failed: {e}", state="error")
            st.stop()

    parsed_m = len(st.session_state.materials or {})
    parsed_p = len(st.session_state.processes or {})
    st.info(f"Parsed **{parsed_m}** materials and **{parsed_p}** processes.")
    if parsed_m == 0:
        st.warning("No materials parsed. Check your columns: Material name/material/name + CO2e + (optional) Recycled/EoL/Lifetime/Circularity.")
        st.stop()
    if parsed_p == 0:
        st.warning("No processes parsed. Ensure the 'Processes' sheet has columns like Process Type + CO2e + Unit (or aliases).")

    st.subheader("Lifetime (Weeks).")
    st.session_state.assessment["lifetime_weeks"] = st.number_input(
        "", min_value=1, value=int(st.session_state.assessment.get("lifetime_weeks", 52))
    )

    st.subheader("Materials & Processes.")
    mats = list(st.session_state.materials.keys())
    st.session_state.assessment["selected_materials"] = st.multiselect(
        "Select Materials.", options=mats,
        default=st.session_state.assessment.get("selected_materials", [])
    )

    if not st.session_state.assessment["selected_materials"]:
        st.info("Select at least one material to proceed.")
        st.stop()

    for m in st.session_state.assessment["selected_materials"]:
        st.markdown(f"### {m}.")
        masses = st.session_state.assessment.setdefault("material_masses", {})
        procs_data = st.session_state.assessment.setdefault("processing_data", {})

        mass_default = float(masses.get(m, 1.0))
        masses[m] = st.number_input(f"Mass Of {m} (kg).", min_value=0.0, value=mass_default, key=f"mass_{m}")

        props = st.session_state.materials[m]
        st.caption(f"COâ‚‚e/kg: {props['COâ‚‚e (kg)']} Â· Recycled %: {props['Recycled Content']} Â· EoL: {props['EoL']}")

        steps = procs_data.setdefault(m, [])
        n = st.number_input(f"How Many Processing Steps For {m}?", min_value=0, max_value=10, value=len(steps), key=f"steps_{m}")
        if n < len(steps):
            steps[:] = steps[:int(n)]
        else:
            for _ in range(int(n) - len(steps)):
                steps.append({"process": "", "amount": 1.0, "co2e_per_unit": 0.0, "unit": ""})

        for i in range(int(n)):
            proc_options = [''] + list(st.session_state.processes.keys())
            current_proc = steps[i]['process'] if steps[i]['process'] in st.session_state.processes else ''
            idx = proc_options.index(current_proc) if current_proc in proc_options else 0
            proc = st.selectbox(
                f"Process #{i+1}.", options=proc_options, index=idx, key=f"proc_{m}_{i}"
            )
            if proc:
                pr = st.session_state.processes.get(proc, {})
                amt = st.number_input(
                    f"Amount For '{proc}' ({pr.get('Unit','')}).",
                    min_value=0.0, value=float(steps[i].get('amount', 1.0)), key=f"amt_{m}_{i}"
                )
                steps[i] = {"process": proc, "amount": amt, "co2e_per_unit": pr.get('COâ‚‚e', 0.0), "unit": pr.get('Unit', '')}

    # normalize structures after user inputs
    _ensure_assessment_model()

# -----------------------------
# Compute results
# -----------------------------

def compute_results():
    data = st.session_state.assessment
    mats = st.session_state.materials
    total_material = 0.0
    total_process  = 0.0
    total_mass     = 0.0
    weighted       = 0.0
    eol            = {}
    cmp_rows       = []
    circ_map = {"high": 3, "medium": 2, "low": 1, "not circular": 0}
    for name in data.get('selected_materials', []):
        m = mats.get(name, {})
        mass = float(data.get('material_masses', {}).get(name, 0))
        total_mass += mass
        total_material += mass * float(m.get('COâ‚‚e (kg)', 0))
        weighted += mass * float(m.get('Recycled Content', 0))
        eol[name] = m.get('EoL', 'Unknown')
        for s in data.get('processing_data', {}).get(name, []):
            total_process += float(s.get('amount', 0)) * float(s.get('co2e_per_unit', 0))
        cmp_rows.append({
            'Material': name,
            'CO2e per kg': float(m.get('COâ‚‚e (kg)', 0)),
            'Recycled Content (%)': float(m.get('Recycled Content', 0)),
            'Circularity (mapped)': circ_map.get(str(m.get('Circularity','')).strip().lower(), 0),
            'Circularity (text)': m.get('Circularity', 'Unknown'),
            'Lifetime (years)': extract_number(m.get('Lifetime', 0)),
            'Lifetime (text)': m.get('Lifetime', 'Unknown'),
        })
    overall = total_material + total_process
    years = max(data.get('lifetime_weeks', 52) / 52, 1e-9)
    return {
        'total_material_co2': total_material,
        'total_process_co2': total_process,
        'overall_co2': overall,
        'weighted_recycled': (weighted / total_mass if total_mass > 0 else 0.0),
        'trees_equiv': overall / (22 * years),
        'total_trees_equiv': overall / 22,
        'lifetime_years': years,
        'eol_summary': eol,
        'comparison': cmp_rows
    }

# -----------------------------
# Report builders (PDF/DOCX)
# -----------------------------

def _material_rows_for_report(selected_materials: List[str], materials_dict: dict, material_masses: dict, lifetime_years: float):
    rows = []
    for m in selected_materials:
        props = materials_dict.get(m, {})
        mass = float(material_masses.get(m, 0.0))
        co2_per_kg = float(props.get("COâ‚‚e (kg)", 0.0))
        co2_total = mass * co2_per_kg
        trees_mat = (co2_total / (22.0 * max(lifetime_years, 1e-9))) if lifetime_years else 0.0
        rows.append([
            m, f"{co2_total:.2f}", f"{float(props.get('Recycled Content', 0.0)):.0f}%",
            str(props.get("Circularity", "Unknown")), str(props.get("EoL", "Unknown")), f"{trees_mat:.1f}"
        ])
    return rows


def build_pdf_from_template(project: str, notes: str, summary: dict, selected_materials: List[str], materials_dict: dict, material_masses: dict) -> Optional[bytes]:
    if not REPORTLAB_OK: return None
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    H1 = styles["Heading1"]; H1.fontSize = 18
    H2 = styles["Heading2"]; H2.fontSize = 14
    P  = styles["BodyText"]; P.leading = 15
    story = []
    if _logo_bytes:
        try:
            story += [RLImage(BytesIO(_logo_bytes), width=120, height=40), Spacer(1, 6)]
        except Exception:
            pass
    story += [Paragraph(f"{project} â€” Easy LCA Report", H1), Spacer(1, 6)]
    story += [Paragraph("Introduction", H2),
              Paragraph("At Tchai we build different: within every brand space we design we try to leave a positive mark on people and planet. Our Easy LCA tool helps us see the real footprint of a concept before itâ€™s built.", P), Spacer(1, 8)]
    story += [Paragraph("Key Metrics", H2)]
    story += [Paragraph(f"Lifetime: <b>{summary['lifetime_years']:.1f} years</b> ({int(summary['lifetime_years']*52)} weeks)", P)]
    story += [Paragraph(f"Total COâ‚‚e: <b>{summary['overall_co2']:.1f} kg</b>", P)]
    story += [Paragraph(f"Weighted recycled content: <b>{summary['weighted_recycled']:.1f}%</b>", P)]
    story += [Paragraph(f"Trees/year: <b>{summary['trees_equiv']:.1f}</b> Â· Total trees: <b>{summary['total_trees_equiv']:.1f}</b>", P)]
    story += [Paragraph("<i>Tree Equivalent is a communication proxy: the estimated number of trees needed to sequester the same COâ‚‚e over your chosen lifetime (assumes ~22 kg COâ‚‚ per tree per year).</i>", P), Spacer(1, 8)]
    if notes: story += [Paragraph("Executive Notes", H2), Paragraph(notes, P), Spacer(1, 8)]
    story += [Paragraph("Material Comparison Overview", H2)]
    header = ["Material", "COâ‚‚e per Unit (kg COâ‚‚e)", "Avg. Recycled Content", "Circularity", "End-of-Life", "Tree Equivalent*"]
    body = _material_rows_for_report(selected_materials, materials_dict, material_masses, summary["lifetime_years"])
    table = Table([header] + body, colWidths=[90, 90, 90, 80, 90, 80])
    table.setStyle(TableStyle([["GRID", (0,0), (-1,-1), 0.6, colors.grey],["BACKGROUND", (0,0), (-1,0), colors.HexColor("#F3F4F6")],["FONTNAME", (0,0), (-1,0), "Helvetica-Bold"],["ALIGN", (1,1), (-1,-1), "CENTER"],["VALIGN", (0,0), (-1,-1), "MIDDLE"]]))
    story += [table, Spacer(1, 6)]
    story += [Paragraph("*Estimated number of trees required to sequester the COâ‚‚e emissions from one unit over the selected years.", P)]
    story += [Spacer(1, 6), Paragraph("End-of-Life Summary", H2)]

    if summary["eol_summary"]:
        bullets = "".join([f"â€¢ <b>{k}</b>: {v}<br/>" for k, v in summary["eol_summary"].items()])
        story += [Paragraph(bullets, P)]
    else:
        story += [Paragraph("â€”", P)]

    story += [
        Spacer(1, 8),
        Paragraph("Conclusion", H2),
        Paragraph("Not every improvement appears in a COâ‚‚e score. Each option presents different strengths and trade-offs.", P)
    ]
    doc.build(story)
    return buf.getvalue()

# ---------- DOCX helpers (step 16) ----------
def _replace_text_in_docx(doc, mapping: dict):
    def replace_in_para(para):
        full_text = "".join(run.text for run in para.runs)
        changed = False
        for k, v in mapping.items():
            if k in full_text:
                full_text = full_text.replace(k, v)
                changed = True
        if changed:
            # clear runs and rebuild one run with merged text
            for _ in range(len(para.runs)):
                r = para.runs[0]
                r.text = ""
                try:
                    r.clear()
                except Exception:
                    pass
            try:
                para.clear()
            except Exception:
                pass
            para.add_run(full_text)

    # paragraphs
    for para in doc.paragraphs:
        replace_in_para(para)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

def _materials_table_block(doc, rows: list):
    doc.add_heading("Material Comparison Overview", level=2)
    headers = ["Material", "COâ‚‚e per Unit (kg COâ‚‚e)", "Avg. Recycled Content", "Circularity", "End-of-Life", "Tree Equivalent*"]
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r in rows:
        c = table.add_row().cells
        for i, v in enumerate(r):
            c[i].text = str(v)
    doc.add_paragraph("*Estimated number of trees required to sequester the COâ‚‚e emissions from one unit over the selected years.")

REPORT_TEMPLATE_PATH = GUIDES / "report_template_cleaned.docx"

def find_report_template() -> Optional[Path]:
    p = REPORT_TEMPLATE_PATH
    try:
        if p.exists() and p.is_file() and p.suffix.lower() == ".docx" and p.stat().st_size > 0:
            # .docx is a ZIP container â€” test integrity
            with zipfile.ZipFile(p, "r") as z:
                z.testzip()
            return p
    except zipfile.BadZipFile:
        return None
    except Exception:
        logger.exception("Template check failed")
        return None
    return None

def build_docx_from_attached_template(project: str, notes: str, R: dict,
                                      selected_materials: List[str], materials_dict: dict, material_masses: dict) -> Optional[bytes]:
    if not DOCX_OK:
        return None
    template = find_report_template()
    if not template:
        return None
    try:
        doc = Document(str(template))
        mapping = {
            "{PROJECT}": project,
            "{LIFETIME_YEARS}": f"{R['lifetime_years']:.1f}",
            "{LIFETIME_WEEKS}": f"{int(R['lifetime_years']*52)}",
            "{TOTAL_CO2}": f"{R['overall_co2']:.1f}",
            "{WEIGHTED_RECYCLED}": f"{R['weighted_recycled']:.1f}%",
            "{TREES_YEAR}": f"{R['trees_equiv']:.1f}",
            "{TREES_TOTAL}": f"{R['total_trees_equiv']:.1f}",
            "{EXEC_NOTES}": (notes or "").strip(),
        }
        _replace_text_in_docx(doc, mapping)
        mat_rows = _material_rows_for_report(selected_materials, materials_dict, material_masses, R["lifetime_years"])
        _materials_table_block(doc, mat_rows)
        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()
    except Exception:
        logger.exception("DOCX-from-template build failed")
        return None

def build_docx_fallback(project: str, notes: str, summary: dict,
                        selected_materials: List[str], materials_dict: dict, material_masses: dict) -> Optional[bytes]:
    if not DOCX_OK:
        return None
    try:
        doc = Document()
        doc.add_heading(f"{project} â€” Easy LCA Report", 0)
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("At Tchai we build differentâ€¦ Our Easy LCA tool helps us see the real footprint of a concept before itâ€™s built.")
        doc.add_heading("Key Metrics", level=1)
        doc.add_paragraph(f"Lifetime: {summary['lifetime_years']:.1f} years ({int(summary['lifetime_years']*52)} weeks)")
        doc.add_paragraph(f"Total COâ‚‚e: {summary['overall_co2']:.1f} kg")
        doc.add_paragraph(f"Weighted recycled content: {summary['weighted_recycled']:.1f}%")
        doc.add_paragraph(f"Trees/year: {summary['trees_equiv']:.1f} Â· Total trees: {summary['total_trees_equiv']:.1f}")
        if notes:
            doc.add_heading("Executive Notes", level=2)
            doc.add_paragraph(notes)
        doc.add_heading("Material Comparison Overview", level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Material"; hdr[1].text = "COâ‚‚e per Unit"; hdr[2].text = "Avg. Recycled Content"; hdr[3].text = "Circularity"; hdr[4].text = "End-of-Life"; hdr[5].text = "Tree Equivalent*"
        for row in _material_rows_for_report(selected_materials, materials_dict, material_masses, summary["lifetime_years"]):
            r = table.add_row().cells
            for i, v in enumerate(row):
                r[i].text = str(v)
        doc.add_heading("End-of-Life Summary", level=1)
        if summary['eol_summary']:
            for k, v in summary['eol_summary'].items():
                doc.add_paragraph(f"â€¢ {k}: {v}")
        else:
            doc.add_paragraph("â€”")
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph("Use these insights to shape a smarter, more sustainable design.")
        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()
    except Exception:
        logger.exception("DOCX-fallback build failed")
        return None

# -----------------------------
# Results (workspace) page
# -----------------------------
if page in (t("nav.results","Results"), "Workspace"):
    if not st.session_state.assessment.get('selected_materials'):
        st.info("Go to Actual Tool and add at least one material.")
        st.stop()

    R = compute_results()
    tab_labels = ["Results & Comparison", "Final Summary", "Report"]
    t0, t1, t2 = st.tabs(tab_labels)

    with t0:
        c1, c2, c3 = st.columns(3)
        c1.metric("Total COâ‚‚ (Materials)", f"{R['total_material_co2']:.1f} kg", help="Embodied carbon from materials only")
        c2.metric("Total COâ‚‚ (Processes)", f"{R['total_process_co2']:.1f} kg", help="Sum of process steps: amount Ã— factor")
        c3.metric("Recycled Content", f"{R['weighted_recycled']:.1f}%", help="Mass-weighted % of recycled inputs")

        df = pd.DataFrame(R['comparison'])
        if df.empty:
            st.info("No data yet.")
        else:
            def style(fig):
                fig.update_layout(
                    plot_bgcolor=BG, paper_bgcolor=BG,
                    font=dict(color="#000", size=14),
                    title_x=0.5, title_font_size=20
                )
                return fig
            a, b = st.columns(2)
            with a:
                fig = px.bar(df, x="Material", y="CO2e per kg", color="Material", title="COâ‚‚e Per Kg", color_discrete_sequence=[POP])
                st.plotly_chart(style(fig), use_container_width=True)
            with b:
                fig = px.bar(df, x="Material", y="Recycled Content (%)", color="Material", title="Recycled Content (%)", color_discrete_sequence=[POP])
                st.plotly_chart(style(fig), use_container_width=True)
            c, d = st.columns(2)
            with c:
                fig = px.bar(df, x="Material", y="Circularity (mapped)", color="Material", title="Circularity", color_discrete_sequence=[POP])
                fig.update_yaxes(tickmode='array', tickvals=[0,1,2,3], ticktext=['Not Circular','Low','Medium','High'])
                st.plotly_chart(style(fig), use_container_width=True)
            with d:
                g = df.copy()
                def life_cat(x):
                    v = extract_number(x)
                    return 'Short' if v < 5 else ('Medium' if v <= 15 else 'Long')
                g['Lifetime Category'] = g['Lifetime (years)'].apply(life_cat)
                MAP = {"Short":1, "Medium":2, "Long":3}
                g['Lifetime'] = g['Lifetime Category'].map(MAP)
                fig = px.bar(g, x="Material", y="Lifetime", color="Material", title="Lifetime", color_discrete_sequence=[POP])
                fig.update_yaxes(tickmode='array', tickvals=[1,2,3], ticktext=['Short','Medium','Long'])
                st.plotly_chart(style(fig), use_container_width=True)

    with t1:
        m1, m2, m3 = st.columns(3)
        m1.markdown(f"<div class='metric'><div>Total Impact COâ‚‚e.</div><h2>{R['overall_co2']:.1f} kg</h2></div>", unsafe_allow_html=True)
        m2.markdown(f"<div class='metric'><div>Tree Equivalent / Year.</div><h2>{R['trees_equiv']:.1f}</h2></div>", unsafe_allow_html=True)
        m3.markdown(f"<div class='metric'><div>Total Trees.</div><h2>{R['total_trees_equiv']:.1f}</h2></div>", unsafe_allow_html=True)
        st.markdown("<p style='margin-top:8px; font-size:0.95rem; color:#374151'><b>Tree Equivalent</b> is a communication proxy (assumes ~22 kg COâ‚‚ per tree per year).</p>", unsafe_allow_html=True)

        st.markdown("#### End-Of-Life Summary.")
        if R['eol_summary']:
            for k, v in R['eol_summary'].items():
                st.write(f"â€¢ **{k}** â€” {v}")
        else:
            st.write("â€”")

    with t2:
        project = st.text_input("Project Name", value="Sample Project")
        notes = st.text_area("Executive Notes")

        templ_docx = None
        if DOCX_OK:
            try:
                templ_docx = build_docx_from_attached_template(
                    project=project, notes=notes, R=R,
                    selected_materials=st.session_state.assessment["selected_materials"],
                    materials_dict=st.session_state.materials,
                    material_masses=st.session_state.assessment["material_masses"],
                )
            except Exception:
                logger.exception("DOCX-from-template failed")
                templ_docx = None

        if templ_docx:
            st.success("Using attached DOCX template with live numbers.")
            st.download_button(
                "â¬‡ï¸ Download Report (DOCX From Template)",
                data=templ_docx,
                file_name=f"TCHAI_Report_{project.replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            pdf_bytes = None
            try:
                pdf_bytes = build_pdf_from_template(
                    project=project, notes=notes,
                    summary={
                        "lifetime_years": R["lifetime_years"],
                        "overall_co2": R["overall_co2"],
                        "weighted_recycled": R["weighted_recycled"],
                        "trees_equiv": R["trees_equiv"],
                        "total_trees_equiv": R["total_trees_equiv"],
                        "eol_summary": R["eol_summary"],
                    },
                    selected_materials=st.session_state.assessment["selected_materials"],
                    materials_dict=st.session_state.materials,
                    material_masses=st.session_state.assessment["material_masses"],
                )
            except Exception:
                logger.exception("PDF report build failed")
                pdf_bytes = None

            if pdf_bytes:
                st.download_button(
                    "â¬‡ï¸ Download PDF Report (Smart-Filled)",
                    data=pdf_bytes,
                    file_name=f"TCHAI_Report_{project.replace(' ','_')}.pdf",
                    mime="application/pdf"
                )
            else:
                st.warning("PDF backend not found (ReportLab). Trying DOCX fallbackâ€¦")
                try:
                    docx_bytes = build_docx_fallback(
                        project, notes,
                        summary={
                            "lifetime_years": R["lifetime_years"],
                            "overall_co2": R["overall_co2"],
                            "weighted_recycled": R["weighted_recycled"],
                            "trees_equiv": R["trees_equiv"],
                            "total_trees_equiv": R["total_trees_equiv"],
                            "eol_summary": R["eol_summary"],
                        },
                        selected_materials=st.session_state.assessment["selected_materials"],
                        materials_dict=st.session_state.materials,
                        material_masses=st.session_state.assessment["material_masses"],
                    )
                except Exception:
                    logger.exception("DOCX fallback failed")
                    docx_bytes = None

                if docx_bytes:
                    st.download_button(
                        "â¬‡ï¸ Download DOCX Report (Smart-Filled)",
                        data=docx_bytes,
                        file_name=f"TCHAI_Report_{project.replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("Neither PDF nor DOCX export is available. Providing a plain-text report.")
                    lines = []
                    title = f"{project} â€” Easy LCA Report"
                    lines.append(title); lines.append("="*len(title)); lines.append("")
                    lines.append("Key Metrics")
                    lines.append(f"- Lifetime: {R['lifetime_years']:.1f} years ({int(R['lifetime_years']*52)} weeks)")
                    lines.append(f"- Total COâ‚‚e: {R['overall_co2']:.1f} kg")
                    lines.append(f"- Weighted recycled content: {R['weighted_recycled']:.1f}%")
                    lines.append(f"- Trees/year: {R['trees_equiv']:.1f}")
                    lines.append(f"- Total trees: {R['total_trees_equiv']:.1f}")
                    plain_txt = "\n".join(lines).encode("utf-8")
                    st.download_button(
                        "â¬‡ï¸ Download Plain-Text Report",
                        data=plain_txt,
                        file_name=f"TCHAI_Report_{project.replace(' ','_')}.txt",
                        mime="text/plain"
                    )

# -----------------------------
# Versions page (step 18)
# -----------------------------
if page in (t("nav.versions","Version"), "ðŸ“ Versions"):
    st.subheader(" Version Management")

    class VM:
        def __init__(self, storage_dir: str = "lca_versions"):
            self.dir = Path(storage_dir); ensure_dir(self.dir)
            self.meta = self.dir / "lca_versions_metadata.json"

        def _load(self):
            return json.loads(self.meta.read_text()) if self.meta.exists() else {}

        def _save(self, m):
            self.meta.write_text(json.dumps(m, indent=2))

        def save(self, name, data, desc=""):
            SAFE_NAME = re.compile(r"^[A-Za-z0-9._ -]{1,64}$")
            m = self._load()
            name = (name or "").strip()
            if not name:
                return False, "Enter a name."
            if not SAFE_NAME.match(name):
                return False, "Only letters/numbers/space/dot/dash/underscore (max 64)."
            if name in m:
                return False, "Name exists."
            fp = (self.dir / f"{name}.json").resolve()
            if self.dir.resolve() not in fp.parents:
                return False, "Invalid name."
            payload = {
                "assessment_data": data,
                "timestamp": datetime.now().isoformat(),
                "description": desc
            }
            fp.write_text(json.dumps(payload))
            m[name] = {
                "filename": fp.name,
                "description": desc,
                "created_at": datetime.now().isoformat(),
                "materials_count": len(data.get('selected_materials', [])),
                "total_co2": data.get('overall_co2', 0)
            }
            self._save(m)
            return True, "Saved"

        def list(self):
            return self._load()

        def load(self, name):
            m = self._load()
            if name not in m:
                return None, "Not found"
            fp = self.dir / m[name]["filename"]
            if not fp.exists():
                return None, "File missing"
            try:
                payload = json.loads(fp.read_text())
                return payload.get("assessment_data", {}), "Loaded"
            except Exception as e:
                return None, f"Read error: {e}"

        def delete(self, name):
            m = self._load()
            if name not in m:
                return False, "Not found"
            fp = self.dir / m[name]["filename"]
            if fp.exists():
                fp.unlink()
            del m[name]
            self._save(m)
            return True, "Deleted"

    if "vm" not in st.session_state:
        st.session_state.vm = VM()
    vm = st.session_state.vm

    t1, t2, t3 = st.tabs(["Save", "Load", "Manage"])

    with t1:
        name = st.text_input("Version Name")
        desc = st.text_area("Description (Optional)")
        if st.button("ðŸ’¾ Save"):
            data = {**st.session_state.assessment}
            data.update(compute_results())
            ok, msg = vm.save(name, data, desc)
            st.success(msg) if ok else st.error(msg)

    with t2:
        meta = vm.list()
        if not meta:
            st.info("No versions saved yet")
        else:
            # small preview table
            preview_rows = []
            for n, info in meta.items():
                preview_rows.append({
                    "Name": n,
                    "Created": info.get("created_at",""),
                    "Materials": info.get("materials_count",0),
                    "Total COâ‚‚": info.get("total_co2",0)
                })
            st.dataframe(pd.DataFrame(preview_rows).sort_values("Created", ascending=False), use_container_width=True)

            sel = st.selectbox("Select Version", list(meta.keys()))
            if st.button("ðŸ“‚ Load"):
                data, msg = vm.load(sel)
                if data:
                    st.session_state.assessment = data
                    _ensure_assessment_model()
                    st.success(msg)
                    st.info("Go to Results to see loaded data.")
                else:
                    st.error(msg)

    with t3:
        meta = vm.list()
        if not meta:
            st.info("Nothing to manage yet")
        else:
            sel = st.selectbox("Select Version To Delete", list(meta.keys()))
            if st.button("ðŸ—‘ï¸ Delete"):
                ok, msg = vm.delete(sel)
                st.success(msg) if ok else st.error(msg)


