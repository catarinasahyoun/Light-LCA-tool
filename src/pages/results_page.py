import io
import re
import json
import io
import textwrap
import pandas as pd
import streamlit as st
import plotly.express as px  # for charts

class ResultsPage:
    @staticmethod
    def _safe_slug(name: str) -> str:
        name = (name or "").strip().replace(" ", "_")
        return re.sub(r"[^A-Za-z0-9._-]", "_", name) or "Unnamed_Project"

    @staticmethod
    def render():
        """Public entry point used by app.py"""
        # Tabs in your requested order/titles
        tab_comp, tab_summary, tab_report = st.tabs([
            "📊 Comparison & Visualizations",
            "🧾 Results Summary",
            "📄 Report",
        ])

        with tab_comp:
            ResultsPage._render_charts_section()

        with tab_summary:
            ResultsPage._render_summary_section()

        with tab_report:
            ResultsPage._render_report_section(R=None)

    # ---------- 1) COMPARISON & VISUALIZATIONS (first tab) ----------
    @staticmethod
    def _render_charts_section():
        st.markdown("### Comparison & Visualizations")

        # Expect list[dict] like in your original app
        comparison_data = st.session_state.get("comparison_data", [])
        if not comparison_data:
            ResultsPage._show_missing_hint(["comparison_data"])
            return

        df_compare = pd.DataFrame(comparison_data)

        # Ensure expected columns exist
        expected_cols = {
            "Material",
            "CO2e per kg",
            "Recycled Content (%)",
            "Circularity (mapped)",
            "Lifetime (years)",
        }
        missing_cols = [c for c in expected_cols if c not in df_compare.columns]
        if missing_cols:
            st.warning(
                "The comparison dataset is missing columns: "
                + ", ".join(missing_cols)
                + ". Please check the Tool page logic that builds `comparison_data`."
            )

        # Helper: lifetime category
        def lifetime_category(lifetime_value):
            try:
                v = float(lifetime_value)
            except Exception:
                v = 0.0
            if v < 5:
                return "Short"
            elif v <= 15:
                return "Medium"
            else:
                return "Long"

        my_color_sequence = ['#2E7D32', '#388E3C', '#4CAF50', '#66BB6A', '#81C784']

        # Two rows of charts, like before
        col1, col2 = st.columns(2)

        # (A) CO2e per kg
        if {"Material", "CO2e per kg"}.issubset(df_compare.columns):
            with col1:
                fig_co2 = px.bar(
                    df_compare, x="Material", y="CO2e per kg",
                    color="Material", title="🏭 CO₂e per kg",
                    color_discrete_sequence=my_color_sequence
                )
                fig_co2.update_layout(
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='#2E7D32'),
                    title_font_size=18,
                    title_x=0.5
                )
                st.plotly_chart(fig_co2, use_container_width=True)
        else:
            with col1:
                st.info("Missing columns for CO₂e chart (need: Material, CO2e per kg).")

        # (B) Recycled Content
        if {"Material", "Recycled Content (%)"}.issubset(df_compare.columns):
            with col2:
                fig_recycled = px.bar(
                    df_compare, x="Material", y="Recycled Content (%)",
                    color="Material", title="♻️ Recycled Content ",
                    color_discrete_sequence=my_color_sequence
                )
                fig_recycled.update_layout(
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='#2E7D32'),
                    title_font_size=18,
                    title_x=0.5
                )
                st.plotly_chart(fig_recycled, use_container_width=True)
        else:
            with col2:
                st.info("Missing columns for Recycled Content chart (need: Material, Recycled Content (%)).")

        col3, col4 = st.columns(2)

        # (C) Circularity
        if {"Material", "Circularity (mapped)"}.issubset(df_compare.columns):
            with col3:
                fig_circularity = px.bar(
                    df_compare, x="Material", y="Circularity (mapped)",
                    color="Material", title="🔄 Circularity ",
                    color_discrete_sequence=my_color_sequence
                )
                fig_circularity.update_layout(
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='#2E7D32'),
                    title_font_size=18,
                    title_x=0.5,
                    yaxis=dict(
                        tickmode='array',
                        tickvals=[0, 1, 2, 3],
                        ticktext=['Not Circular', 'Low', 'Medium', 'High']
                    )
                )
                st.plotly_chart(fig_circularity, use_container_width=True)
        else:
            with col3:
                st.info("Missing columns for Circularity chart (need: Material, Circularity (mapped)).")

        # (D) Lifetime (Short/Medium/Long)
        if "Lifetime (years)" in df_compare.columns and "Material" in df_compare.columns:
            df_life = df_compare.copy()
            df_life["Lifetime Category"] = df_life["Lifetime (years)"].apply(lifetime_category)
            lifetime_cat_to_num = {"Short": 1, "Medium": 2, "Long": 3}
            df_life["Lifetime"] = df_life["Lifetime Category"].map(lifetime_cat_to_num)

            with col4:
                fig_lifetime = px.bar(
                    df_life, x="Material", y="Lifetime",
                    color="Material", title="⏱️ Lifetime ",
                    color_discrete_sequence=my_color_sequence
                )
                fig_lifetime.update_layout(
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='#2E7D32'),
                    title_font_size=18,
                    title_x=0.5,
                    yaxis=dict(
                        tickmode='array',
                        tickvals=[1, 2, 3],
                        ticktext=["Short", "Medium", "Long"]
                    )
                )
                st.plotly_chart(fig_lifetime, use_container_width=True)
        else:
            with col4:
                st.info("Missing column for Lifetime chart (need: Lifetime (years)).")

    # ---------- 2) RESULTS SUMMARY (second tab) ----------
    @staticmethod
    def _render_summary_section():
        # ======= Styles for boxed KPIs =======
        st.markdown(
            """
            <style>
              .kpi-grid { display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 12px; }
              @media(max-width: 1024px){ .kpi-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); } }
              @media(max-width: 640px){ .kpi-grid { grid-template-columns: 1fr; } }
              .kpi-card {
                border: 1px solid #e0e0e0; border-radius: 12px; padding: 16px;
                background: #ffffff; box-shadow: 0 1px 2px rgba(0,0,0,0.04);
              }
              .kpi-title { font-size: 13px; color: #2E7D32; margin: 0 0 8px 0; font-weight: 600; }
              .kpi-value { font-size: 22px; margin: 0; font-weight: 700; }
              .kpi-sub { font-size: 12px; color: #6b7280; margin-top: 6px; }
            </style>
            """,
            unsafe_allow_html=True,
        )

        st.markdown("### Results Summary")

        # ---- Pull the numbers from session (or fallbacks) ----
        total_material_co2 = float(st.session_state.get("total_material_co2", 0.0) or 0.0)
        total_process_co2  = float(st.session_state.get("total_process_co2", 0.0) or 0.0)
        overall_co2        = float(st.session_state.get("overall_co2", total_material_co2 + total_process_co2) or 0.0)
        weighted_recycled  = float(st.session_state.get("weighted_recycled", 0.0) or 0.0)

        lifetime_weeks     = int(st.session_state.get("lifetime_weeks", 52) or 52)
        lifetime_years     = max(lifetime_weeks / 52.0, 1e-9)  # avoid div/zero

        # ---- Tree equivalent logic (no hard-coded 5 years) ----
        TREE_SEQUESTRATION_PER_YEAR = 22.0  # kg CO2 per tree per year (simple signal)

        mode = st.radio(
            "Tree equivalent basis",
            options=("Over lifetime", "Per year"),
            index=0,
            horizontal=True,
            key="trees_mode_select",
            help="Choose whether to express tree equivalent across the design's lifetime or per year.",
        )

        if mode == "Over lifetime":
            trees_equiv = overall_co2 / (TREE_SEQUESTRATION_PER_YEAR * lifetime_years)
            trees_label = f"{trees_equiv:.2f} trees over {lifetime_years:.1f} years"
            trees_sub   = f"22 kg CO₂/tree/year · lifetime={lifetime_years:.1f} years"
        else:
            trees_equiv = overall_co2 / TREE_SEQUESTRATION_PER_YEAR
            trees_label = f"{trees_equiv:.2f} trees / year"
            trees_sub   = "22 kg CO₂/tree/year"

        # ---- KPI Grid (boxed visuals) ----
        st.markdown('<div class="kpi-grid">', unsafe_allow_html=True)

        # Weighted Recycled
        st.markdown(
            f"""
            <div class="kpi-card">
              <p class="kpi-title">Weighted Recycled Content</p>
              <p class="kpi-value">{weighted_recycled:.1f}%</p>
              <div class="kpi-sub">Mass-weighted across all selected materials</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Total CO2 - Materials
        st.markdown(
            f"""
            <div class="kpi-card">
              <p class="kpi-title">Total CO₂ — Materials</p>
              <p class="kpi-value">{total_material_co2:.2f} kg</p>
              <div class="kpi-sub">Sum of (mass × CO₂e/kg) per material</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Total CO2 - Processes
        st.markdown(
            f"""
            <div class="kpi-card">
              <p class="kpi-title">Total CO₂ — Processes</p>
              <p class="kpi-value">{total_process_co2:.2f} kg</p>
              <div class="kpi-sub">Sum of (amount × CO₂e/unit) for all steps</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Overall CO2
        st.markdown(
            f"""
            <div class="kpi-card">
              <p class="kpi-title">Overall CO₂</p>
              <p class="kpi-value">{overall_co2:.2f} kg</p>
              <div class="kpi-sub">Materials + Processes</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Tree Equivalent (mode-aware)
        st.markdown(
            f"""
            <div class="kpi-card">
              <p class="kpi-title">Tree Equivalent</p>
              <p class="kpi-value">{trees_label}</p>
              <div class="kpi-sub">{trees_sub}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Lifetime (display signal)
        st.markdown(
            f"""
            <div class="kpi-card">
              <p class="kpi-title">Lifetime</p>
              <p class="kpi-value">{lifetime_weeks} weeks</p>
              <div class="kpi-sub">{lifetime_years:.1f} years</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown('</div>', unsafe_allow_html=True)

    # ---------- 3) REPORT (third tab; DOCX/PDF only, custom title, TCHAI logo) ----------
    @staticmethod
    def _render_report_section(R):
        """
        Export a report that follows the Tchai template:
        - Cover with logo + title + date
        - Introduction (from text brief)
        - Golden Rules & Must-Haves
        - Methodology
        - Results KPIs
        - Material Comparison Overview
        - Conclusion
        Formats: DOCX (always) and PDF (if reportlab is installed)
        """
        import re

        # --- deps ---
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except Exception:
            st.error("Missing dependency: install `python-docx` to export DOCX reports.")
            return

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            REPORTLAB_AVAILABLE = True
        except Exception:
            REPORTLAB_AVAILABLE = False

        st.markdown("### Report")

        # ---------- static text pulled from your docs (curated & concise) ----------
        INTRO_TEXT = (
            "At Tchai we build different. Our Easy LCA tool helps us see the footprint of a concept "
            "before it’s built—so we can adjust, swap, or simplify. This report shares early-stage, "
            "directional indicators (materials, processes, end-of-life) to support better decisions."
        )

        GOLDEN_RULES = [
            "Think circular — plan reuse, modularity and end-of-life from day one.",
            "Start light — run LCA-Light in briefing/concept to steer choices early.",
            "Choose smart materials — prefer recycled/recyclable, FSC/PEFC.",
            "Cut waste — simplify parts/finishes/packaging; avoid over-engineering.",
            "Design for disassembly — standard fasteners; clear material separation.",
        ]

        MUST_HAVES = {
            "Materials": [
                "FSC/PEFC wood; sustainable MDF where feasible",
                "Water-based paints / low-VOC adhesives",
                "Avoid PVC/problematic plastics where possible",
                "High recycled content where quality allows",
            ],
            "Design & Build": [
                "Modular & repairable assemblies",
                "Avoid mixed-material laminates that block recycling",
                "Label parts for sorting (material codes)",
                "Standardize repeat parts across programs",
            ],
            "Process & Logistics": [
                "Run LCA-Light before locking specs",
                "Transport-efficient (flat-pack, stackable, lighter)",
                "Source locally when it truly reduces impact",
                "Plan end-of-life (reuse/recycle routes documented)",
            ],
        }

        METHODOLOGY_TEXT = (
            "Scope includes materials and processes with a mass-weighted recycled content signal, "
            "plus a qualitative circularity cue. Tree-equivalent is a simple signal using 22 kg CO₂ "
            "sequestration per tree per year. Transport and social criteria are excluded at this stage."
        )

        CONCLUSION_TEXT = (
            "Not every improvement shows up in a CO₂e score. Each option has trade-offs; the win is "
            "combining insights to shape a smarter, more sustainable design. This indicator report "
            "guides early decisions; a full LCA can follow for official claims."
        )

        # ---------- gather data ----------
        project_name = st.session_state.get("project_name") or (
            (R or {}).get("project_name") if isinstance(R, dict) else None
        )
        def _safe_slug(name: str) -> str:
            name = (name or "").strip().replace(" ", "_")
            return re.sub(r"[^A-Za-z0-9._-]", "_", name) or "Unnamed_Project"
        project_slug = _safe_slug(project_name or "Unnamed_Project")

        comparison_data = (
            st.session_state.get("comparison_data")
            or ((R or {}).get("comparison_data") if isinstance(R, dict) else [])
        )
        df_compare = pd.DataFrame(comparison_data) if comparison_data else pd.DataFrame()

        eol_summary = st.session_state.get("eol_summary", {})  # {material: EoL text}
        totals = {
            "total_material_co2": float(st.session_state.get("total_material_co2") or 0.0),
            "total_process_co2":  float(st.session_state.get("total_process_co2") or 0.0),
            "overall_co2":        float(st.session_state.get("overall_co2") or 0.0),
            "weighted_recycled":  float(st.session_state.get("weighted_recycled") or 0.0),
            "lifetime_weeks":     int(st.session_state.get("lifetime_weeks") or 52),
        }
        lifetime_years = max(totals["lifetime_weeks"] / 52.0, 1e-9)

        # ---------- UI controls ----------
        default_title = f"Easy LCA Tool Report — {project_slug}"
        report_title = st.text_input("Report title", value=default_title, key="report_title_input")

        fmt_options = ["DOCX (.docx)"] + (["PDF (.pdf)"] if REPORTLAB_AVAILABLE else [])
        report_choice = st.selectbox("Format", fmt_options, index=0, key="report_format_choice")
        if not REPORTLAB_AVAILABLE:
            st.caption("Install `reportlab` to enable PDF export.")

        # ---------- logo loader ----------
        def _load_logo():
            for p in [
                "/mnt/data/tchai_logo.png",    # provided path
                "assets/logo/tchai_logo.png",
                "assets/tchai_logo.png",
                "tchai_logo.png",
            ]:
                try:
                    with open(p, "rb") as f:
                        return f.read()
                except Exception:
                    continue
            return None
        logo_bytes = _load_logo()

        # ---------- helpers ----------
        def _docx_heading(doc, text, level=1):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14 if level == 1 else 12)
            return p

        def _docx_bullets(doc, items):
            for it in items:
                para = doc.add_paragraph(style=None)
                para_format = para.paragraph_format
                run = para.add_run(f"• {it}")
                run.font.size = Pt(11)

        # ---------- DOCX builder ----------
        def build_docx() -> bytes:
            doc = Document()
            sec = doc.sections[0]
            sec.top_margin = Inches(0.7)
            sec.bottom_margin = Inches(0.7)
            sec.left_margin = Inches(0.7)
            sec.right_margin = Inches(0.7)

            # Cover
            if logo_bytes:
                tmp = io.BytesIO(logo_bytes)
                try:
                    doc.add_picture(tmp, width=Inches(1.6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    pass
            title_p = doc.add_paragraph()
            t = title_p.add_run(report_title)
            t.bold = True
            t.font.size = Pt(20)
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            meta = doc.add_paragraph(f"Project: {project_name or project_slug}   ·   Date: {pd.Timestamp.now():%Y-%m-%d}")
            meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph("")  # spacer

            # Introduction
            _docx_heading(doc, "Introduction", level=1)
            doc.add_paragraph(INTRO_TEXT)

            # Golden Rules & Must-Haves
            _docx_heading(doc, "Tchai Conscious Standard", level=1)
            doc.add_paragraph("5 Golden Rules")
            _docx_bullets(doc, GOLDEN_RULES)
            doc.add_paragraph("")  # spacer
            doc.add_paragraph("12 Must-Haves")
            for section, bullets in MUST_HAVES.items():
                doc.add_paragraph(section).runs[0].bold = True
                _docx_bullets(doc, bullets)

            # Methodology
            _docx_heading(doc, "Methodology", level=1)
            doc.add_paragraph(f"Lifespan set: {totals['lifetime_weeks']} weeks ({lifetime_years:.1f} years).")
            doc.add_paragraph(METHODOLOGY_TEXT)

            # KPIs
            _docx_heading(doc, "Results Summary", level=1)
            tbl = doc.add_table(rows=0, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.style = "Light Grid"
            def add_kpi(label, value):
                row = tbl.add_row().cells
                row[0].text = label
                row[1].text = value
            add_kpi("Weighted Recycled Content", f"{totals['weighted_recycled']:.1f} %")
            add_kpi("Total CO₂ — Materials", f"{totals['total_material_co2']:.2f} kg")
            add_kpi("Total CO₂ — Processes", f"{totals['total_process_co2']:.2f} kg")
            add_kpi("Overall CO₂", f"{totals['overall_co2']:.2f} kg")
            add_kpi("Tree-Equivalent (signal)",
                    f"{(totals['overall_co2']/(22.0*lifetime_years)):.2f} trees over {lifetime_years:.1f} years")

            # Material Comparison Overview
            _docx_heading(doc, "Material Comparison Overview", level=1)
            # Decide columns to show
            cols_pref = ["Material", "CO2e per kg", "Recycled Content (%)", "Circularity (mapped)", "Lifetime (years)"]
            cols_show = [c for c in cols_pref if c in df_compare.columns]
            extra_eol = bool(eol_summary)
            ncols = len(cols_show) + (1 if extra_eol else 0)
            if ncols == 0:
                doc.add_paragraph("No comparison data available.")
            else:
                table = doc.add_table(rows=1, cols=ncols)
                table.style = "Light Grid"
                hdr_cells = table.rows[0].cells
                for j, col in enumerate(cols_show):
                    hdr_cells[j].text = col
                if extra_eol:
                    hdr_cells[len(cols_show)].text = "End-of-Life"
                for _, row in df_compare.iterrows():
                    mat = row.get("Material", "")
                    cells = table.add_row().cells
                    for j, col in enumerate(cols_show):
                        cells[j].text = str(row.get(col, ""))
                    if extra_eol:
                        cells[len(cols_show)].text = str(eol_summary.get(mat, ""))

            # Conclusion
            _docx_heading(doc, "Conclusion", level=1)
            doc.add_paragraph(CONCLUSION_TEXT)

            out = io.BytesIO()
            doc.save(out)
            return out.getvalue()

        # ---------- simple PDF builder (one page summary) ----------
        def build_pdf() -> bytes:
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=A4)
            w, h = A4
            x = 50
            y = h - 50

            def line(text, font="Helvetica", size=10, dy=14, bold=False):
                nonlocal y
                c.setFont("Helvetica-Bold" if bold else font, size)
                c.drawString(x, y, text)
                y -= dy

            # logo + title
            if logo_bytes:
                try:
                    c.drawImage(ImageReader(io.BytesIO(logo_bytes)), x, y-35, width=110, preserveAspectRatio=True, mask='auto')
                except Exception:
                    pass
            line(report_title, size=16, dy=22, bold=True); y -= 6
            line(f"Project: {project_name or project_slug}   ·   Date: {pd.Timestamp.now():%Y-%m-%d}", size=9, dy=16)

            # intro
            line("Introduction", bold=True, dy=18)
            for chunk in text_wrap(INTRO_TEXT, width=92) if hasattr(st, "text_wrap") else [INTRO_TEXT]:
                line(chunk, size=9, dy=12)

            # KPIs
            y -= 6
            line("Results Summary", bold=True, dy=18)
            line(f"Weighted Recycled Content: {totals['weighted_recycled']:.1f} %", size=9)
            line(f"Total CO₂ — Materials: {totals['total_material_co2']:.2f} kg", size=9)
            line(f"Total CO₂ — Processes: {totals['total_process_co2']:.2f} kg", size=9)
            line(f"Overall CO₂: {totals['overall_co2']:.2f} kg", size=9)
            line(f"Tree-Equivalent (signal): {(totals['overall_co2']/(22.0*lifetime_years)):.2f} trees over {lifetime_years:.1f} years", size=9)

            # footer
            c.setFont("Helvetica-Oblique", 8)
            c.drawRightString(w-40, 28, "Generated with TCHAI Easy LCA Tool")
            c.showPage(); c.save()
            return buf.getvalue()

        # ---------- build + download ----------
        if report_choice.startswith("DOCX"):
            data_bytes = build_docx()
            file_name = f"{_safe_slug(report_title)}.docx"
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            if not REPORTLAB_AVAILABLE:
                st.error("PDF export unavailable (install `reportlab`).")
                return
            data_bytes = build_pdf()
            file_name = f"{_safe_slug(report_title)}.pdf"
            mime = "application/pdf"

        st.download_button(
            label=f"⬇️ Download {file_name}",
            data=data_bytes,
            file_name=file_name,
            mime=mime,
            use_container_width=True,
        )
